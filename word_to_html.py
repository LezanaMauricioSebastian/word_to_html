# Author: Lezana Mauricio Sebastian
import tkinter as tk
from tkinter import filedialog, scrolledtext
from docx import Document

# Function to convert relevant Word content to HTML
def word_to_html(docx_file):
    document = Document(docx_file)
    html_content = ""
    capture_data = False
    list_open = False
    html_content += "\n \n"
    cont = 0
    for para in document.paragraphs:
        # Check for the "Features" heading
        if "Features" in para.text or "Features:" in para.text:
            if list_open:  # Close the previous <ul> if open
                html_content += "</ul>\n"
            html_content += "<h2>Features</h2>\n<ul>\n"
            capture_data = True
            list_open = True
        # Check for the "Technical Specifications" heading 
        elif "Technical Specification:" in para.text or "Technical Specifications:" in para.text or "Technical specification:" in para.text:
            if list_open:  # Close the <ul> if open
                html_content += "</ul>\n"
                list_open = False
            capture_data = True
            html_content += "&nbsp;\n"
            html_content += "<h2>Technical Specifications</h2>\n"
            
            # Capture only the next available table if it exists
            if cont < len(document.tables):
                table = document.tables[cont]
                cont += 1
                
                html_content += "<table>\n"
                for row in table.rows:
                    html_content += "<tr>\n"
                    html_content += f"<td>{row.cells[0].text.strip()}</td>\n"
                    html_content += f"<td>{row.cells[1].text.strip()}</td>\n"
                    html_content += "</tr>\n"
                html_content += "</table>\n"
                html_content += "\n \n NEXT PRODUCT  \n \n"
        # Add list items or paragraphs if in "Features"
        elif capture_data and para.text.strip() != "":
            if list_open and (para.text.startswith("•") or para.text.startswith("-")):
                html_content += f"<li>{para.text.strip()}</li>\n"
            elif list_open:
                # Check if it's not a link or unwanted data
                if not any(word in para.text for word in ["Alibaba Link:", "Alibaba", "Supplier Link:", "Prices", "Meta Description","Product Link:"]):
                    html_content += f"<li>{para.text.strip()}</li>\n" 
                else:
                    capture_data = False  # Stop capturing features if unwanted text appears

        # Stop capturing data if unwanted text appears
        if any(word in para.text for word in ["Alibaba Link:", "Alibaba", "Supplier Link:", "Prices", "Meta Description","Product Link:"]):
            capture_data = False

    if list_open:  # Close any unclosed list
        html_content += "</ul>\n"

    return html_content

# Function to open the Word file and convert it
def open_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        html_result = word_to_html(file_path)
        # Display the result in the text box
        text_box.delete(1.0, tk.END)
        text_box.insert(tk.END, html_result)

# Create the main window
root = tk.Tk()
root.title("Word to HTML Converter")

# Button to select the Word file
btn_open = tk.Button(root, text="Select Word File", command=open_file)
btn_open.pack(pady=10)

# Text box to display the resulting HTML
text_box = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=100, height=40)
text_box.pack(padx=10, pady=10, fill=tk.BOTH, expand=True)

# Run the Tkinter application
root.mainloop()

